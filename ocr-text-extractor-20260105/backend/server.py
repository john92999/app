from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Define Models
class StatusCheck(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class StatusCheckCreate(BaseModel):
    client_name: str

class TextToDocxRequest(BaseModel):
    text: str
    filename: str = "extracted_text.docx"

# Routes
@api_router.get("/")
async def root():
    return {"message": "OCR App Backend API"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.dict()
    status_obj = StatusCheck(**status_dict)
    _ = await db.status_checks.insert_one(status_obj.dict())
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    status_checks = await db.status_checks.find().to_list(1000)
    return [StatusCheck(**status_check) for status_check in status_checks]

@api_router.post("/generate-docx")
async def generate_docx(request: TextToDocxRequest):
    """
    Generate a DOCX file from extracted text
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Add a title
        heading = doc.add_heading('Extracted Text', level=1)
        
        # Split text into paragraphs and add them
        paragraphs = request.text.split('\n')
        for para in paragraphs:
            if para.strip():  # Only add non-empty paragraphs
                p = doc.add_paragraph(para)
                # Set font size
                for run in p.runs:
                    run.font.size = Pt(11)
        
        # Save to bytes buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Return as streaming response
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={request.filename}"
            }
        )
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()